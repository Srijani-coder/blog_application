"""Rich DOCX/RTF/TXT importer for StatDash posts.

Converts uploaded Word documents into semantic HTML while preserving as much
formatting as possible: headings, paragraphs, fonts, bold/italic/underline,
colors, hyperlinks, tables, lists, spacing, alignment, and inline images.
Images are uploaded to Cloudinary and placed where they appear in the DOCX.
"""

from __future__ import annotations

import html
import io
import re
from typing import Iterable, List, Optional

import cloudinary.uploader
from docx import Document
from docx.document import Document as _Document
from docx.oxml.ns import qn
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from striprtf.striprtf import rtf_to_text


def _safe(text: object) -> str:
    return html.escape(str(text or ""), quote=True)


def _css_value(value) -> Optional[str]:
    try:
        if value is None:
            return None
        # python-docx Length values expose pt
        if hasattr(value, "pt"):
            return f"{value.pt:.2f}pt"
        return str(value)
    except Exception:
        return None


def _color_to_css(color) -> Optional[str]:
    try:
        if color and color.rgb:
            return f"#{color.rgb}"
    except Exception:
        return None
    return None


def _paragraph_css(paragraph: Paragraph) -> str:
    fmt = paragraph.paragraph_format
    styles = []

    if paragraph.alignment is not None:
        align_map = {0: "left", 1: "center", 2: "right", 3: "justify"}
        val = align_map.get(int(paragraph.alignment))
        if val:
            styles.append(f"text-align:{val}")

    for prop, css in [
        (fmt.left_indent, "margin-left"),
        (fmt.right_indent, "margin-right"),
        (fmt.first_line_indent, "text-indent"),
        (fmt.space_before, "margin-top"),
        (fmt.space_after, "margin-bottom"),
    ]:
        v = _css_value(prop)
        if v:
            styles.append(f"{css}:{v}")

    if fmt.line_spacing:
        try:
            if hasattr(fmt.line_spacing, "pt"):
                styles.append(f"line-height:{fmt.line_spacing.pt:.2f}pt")
            else:
                styles.append(f"line-height:{fmt.line_spacing}")
        except Exception:
            pass

    return ";".join(styles)


def _run_css(run: Run) -> str:
    styles = []
    font = run.font

    if font.name:
        styles.append(f"font-family:'{_safe(font.name)}'")
    if font.size:
        styles.append(f"font-size:{font.size.pt:.2f}pt")
    color = _color_to_css(font.color)
    if color:
        styles.append(f"color:{color}")
    if font.highlight_color:
        # Word highlight names do not map perfectly to CSS. Keep a readable marker.
        styles.append("background-color:yellow")
    if font.all_caps:
        styles.append("text-transform:uppercase")
    if font.small_caps:
        styles.append("font-variant:small-caps")
    if font.strike:
        styles.append("text-decoration:line-through")
    if font.superscript:
        styles.append("vertical-align:super;font-size:smaller")
    if font.subscript:
        styles.append("vertical-align:sub;font-size:smaller")

    return ";".join(styles)


def _apply_inline_tags(text: str, run: Run, extra_css: str = "") -> str:
    css = ";".join([v for v in [_run_css(run), extra_css] if v])
    out = text
    if css:
        out = f'<span style="{css}">{out}</span>'
    if run.bold:
        out = f"<strong>{out}</strong>"
    if run.italic:
        out = f"<em>{out}</em>"
    if run.underline:
        out = f"<u>{out}</u>"
    return out


def _image_html(doc, rel_id: str, alt_text: str, css_class: str = "docx-inline-image") -> str:
    try:
        image_part = doc.part.related_parts[rel_id]
        blob = image_part.blob
        res = cloudinary.uploader.upload(
            io.BytesIO(blob),
            folder="statsdash/content_images",
            resource_type="image",
            transformation=[{"quality": "auto"}, {"fetch_format": "auto"}],
        )
        src = res.get("secure_url") or res.get("url")
        if not src:
            return ""
        return (
            f'<figure class="docx-figure">'
            f'<img src="{_safe(src)}" class="media {css_class}" alt="{_safe(alt_text)}" loading="lazy">'
            f'</figure>'
        )
    except Exception as exc:
        return f'<!-- Image import failed: {_safe(exc)} -->'


def _run_to_html(doc, run: Run, image_alt_texts: List[str], image_counter: List[int]) -> str:
    parts: List[str] = []
    # Preserve tabs, line breaks, spaces and drawings in exact run-child order.
    for child in run._element:
        tag = child.tag
        if tag == qn("w:t"):
            text = child.text or ""
            parts.append(_safe(text).replace("  ", " &nbsp;"))
        elif tag == qn("w:tab"):
            parts.append("&emsp;")
        elif tag in {qn("w:br"), qn("w:cr")}:
            parts.append("<br>")
        elif tag == qn("w:drawing"):
            embeds = child.xpath('.//*[local-name()="blip"]/@*[local-name()="embed"]')
            for rel_id in embeds:
                idx = image_counter[0]
                image_counter[0] += 1
                alt = image_alt_texts[idx] if idx < len(image_alt_texts) and image_alt_texts[idx] else f"Article image {idx + 1}"
                parts.append(_image_html(doc, rel_id, alt))
    return _apply_inline_tags("".join(parts), run)


def _hyperlink_to_html(doc, hyperlink, image_alt_texts: List[str], image_counter: List[int]) -> str:
    r_id = hyperlink.get(qn("r:id"))
    href = "#"
    if r_id and r_id in doc.part.rels:
        href = doc.part.rels[r_id].target_ref
    text_parts = []
    for r_el in hyperlink.findall(qn("w:r")):
        run = Run(r_el, None)
        text_parts.append(_run_to_html(doc, run, image_alt_texts, image_counter))
    return f'<a href="{_safe(href)}" target="_blank" rel="noopener noreferrer">{"".join(text_parts)}</a>'


def _paragraph_inner_html(doc, paragraph: Paragraph, image_alt_texts: List[str], image_counter: List[int]) -> str:
    parts: List[str] = []
    for child in paragraph._element:
        if child.tag == qn("w:r"):
            parts.append(_run_to_html(doc, Run(child, paragraph), image_alt_texts, image_counter))
        elif child.tag == qn("w:hyperlink"):
            parts.append(_hyperlink_to_html(doc, child, image_alt_texts, image_counter))
    return "".join(parts).strip()


def _is_list_paragraph(paragraph: Paragraph) -> bool:
    style_name = (paragraph.style.name if paragraph.style else "").lower()
    if "list" in style_name:
        return True
    p_pr = paragraph._p.pPr
    return bool(p_pr is not None and p_pr.numPr is not None)


def _list_kind(paragraph: Paragraph) -> str:
    style_name = (paragraph.style.name if paragraph.style else "").lower()
    if "number" in style_name:
        return "ol"
    return "ul"


def _paragraph_to_html(doc, paragraph: Paragraph, image_alt_texts: List[str], image_counter: List[int]) -> str:
    inner = _paragraph_inner_html(doc, paragraph, image_alt_texts, image_counter)
    if not inner:
        # preserve blank line/spacing from Word document
        return '<p class="docx-empty-line">&nbsp;</p>'

    style_name = (paragraph.style.name if paragraph.style else "").lower()
    css = _paragraph_css(paragraph)
    style_attr = f' style="{css}"' if css else ""

    if style_name.startswith("heading"):
        match = re.search(r"(\d+)", style_name)
        level = min(6, max(1, int(match.group(1)) if match else 2))
        return f"<h{level}{style_attr}>{inner}</h{level}>"

    if _is_list_paragraph(paragraph):
        return f'<li{style_attr}>{inner}</li>'

    return f'<p class="docx-paragraph"{style_attr}>{inner}</p>'


def _table_to_html(doc, table: Table, image_alt_texts: List[str], image_counter: List[int]) -> str:
    rows = []
    for row in table.rows:
        cells = []
        for cell in row.cells:
            cell_html = []
            for p in cell.paragraphs:
                cell_html.append(_paragraph_to_html(doc, p, image_alt_texts, image_counter))
            cells.append(f"<td>{''.join(cell_html)}</td>")
        rows.append(f"<tr>{''.join(cells)}</tr>")
    return f'<div class="docx-table-wrap"><table class="docx-table">{"".join(rows)}</table></div>'


def _iter_block_items(parent) -> Iterable[Paragraph | Table]:
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        parent_elm = parent.element.body

    for child in parent_elm.iterchildren():
        if child.tag.endswith("}p"):
            yield Paragraph(child, parent)
        elif child.tag.endswith("}tbl"):
            yield Table(child, parent)


def _group_lists(html_parts: List[str], paragraphs: List[Paragraph]) -> str:
    output: List[str] = []
    open_list: Optional[str] = None
    p_index = 0
    for part in html_parts:
        if part.startswith("<li"):
            kind = _list_kind(paragraphs[p_index]) if p_index < len(paragraphs) else "ul"
            if open_list != kind:
                if open_list:
                    output.append(f"</{open_list}>")
                output.append(f"<{kind} class=\"docx-list\">")
                open_list = kind
            output.append(part)
        else:
            if open_list:
                output.append(f"</{open_list}>")
                open_list = None
            output.append(part)
        p_index += 1
    if open_list:
        output.append(f"</{open_list}>")
    return "\n".join(output)


def extract_rich_content(file, image_alt_texts: Optional[List[str]] = None) -> str:
    name = (getattr(file, "filename", "") or "").lower()
    image_alt_texts = image_alt_texts or []

    if name.endswith(".docx"):
        doc = Document(file)
        image_counter = [0]
        output: List[str] = []
        open_list: Optional[str] = None

        for block in _iter_block_items(doc):
            if isinstance(block, Paragraph):
                part = _paragraph_to_html(doc, block, image_alt_texts, image_counter)
                if part.startswith("<li"):
                    kind = _list_kind(block)
                    if open_list != kind:
                        if open_list:
                            output.append(f"</{open_list}>")
                        output.append(f"<{kind} class=\"docx-list\">")
                        open_list = kind
                    output.append(part)
                else:
                    if open_list:
                        output.append(f"</{open_list}>")
                        open_list = None
                    output.append(part)
            elif isinstance(block, Table):
                if open_list:
                    output.append(f"</{open_list}>")
                    open_list = None
                output.append(_table_to_html(doc, block, image_alt_texts, image_counter))

        if open_list:
            output.append(f"</{open_list}>")
        return "\n".join(output)

    raw = file.read()
    if isinstance(raw, bytes):
        text = raw.decode("utf-8", errors="ignore")
    else:
        text = raw

    if name.endswith(".rtf"):
        text = rtf_to_text(text)

    return "\n".join(
        f'<p class="docx-paragraph">{_safe(line).replace("  ", " &nbsp;")}</p>'
        for line in text.splitlines()
        if line.strip()
    )


def count_docx_images(file) -> int:
    """Count images in DOCX. Useful for optional UI/validation workflows."""
    pos = file.stream.tell()
    try:
        file.stream.seek(0)
        doc = Document(file)
        count = 0
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                count += 1
        return count
    finally:
        file.stream.seek(pos)
